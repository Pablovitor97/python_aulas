from docx import Document
from docx.shared import Inches
document = Document()
document.add_heading('Titulo', 0)

document.add_paragraph('A plain paragraph having some ')

document.add_heading('Heading, level 1', level=1)

document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)

document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.save('c:/temp/pabloneruda.docx')

